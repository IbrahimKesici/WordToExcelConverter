from docx import Document


class DocManager():

    def __init__(self, fileName, headers):
        self.fileName = fileName
        self.doc = Document(fileName)
        self.paragraphs = [paragraph for paragraph in self.doc.paragraphs if not paragraph.text.rstrip() in '']
        self.tables = self.doc.tables[:2]
        self.headers = headers

    def getDetails(self):
        countryCode = self.fileName.split("_")[-4]
        year = self.fileName.split("_")[-2]
        
        for i in range(0, len(self.paragraphs)):
            if self.paragraphs[i].text.lower() == "overall evaluation":
                textSplited = self.paragraphs[i+1].text.split(", ")
                cityName = textSplited[0]
                countryName = textSplited[1]
                break
        
        return cityName, countryName, countryCode, year

    def getRating(self):
        ratingTable = self.tables[0]
        scoreTable = self.tables[1]

        ratings = []
        for row in ratingTable.rows[1:]:
            ratings.append(row.cells[1].paragraphs[0].text)

        for row in scoreTable.rows:
            ratings.append(row.cells[1].paragraphs[0].text)

        return ratings

    def getDescription(self, startIndex):
        checkHeader = ["factor ratings","\nfactor ratings"]
        
        for i in range(0, len(self.paragraphs)):
            if self.paragraphs[i].text.lower() in checkHeader:
                index = i + startIndex
                break

        description = []
        for header in self.headers[5:]:
            description.append(self.paragraphs[index].text)
            index += 1

        return description

    